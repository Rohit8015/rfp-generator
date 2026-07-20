"""Document parsing helpers — Phase 3.

Format-dispatching text extraction for markdown, pdf, docx and xlsx.

Everything here is deterministic. The structurer calls this first and only reaches for a
model when a block cannot be labelled from the text itself, so a parsing regression
shows up as a parsing bug rather than as a mysterious drop in extraction quality.
"""

from __future__ import annotations

import re
from pathlib import Path

#: Extensions we can read. Anything else raises rather than silently returning "".
SUPPORTED = {".md", ".markdown", ".txt", ".pdf", ".docx", ".xlsx"}


class UnsupportedDocument(ValueError):
    """The file extension has no parser."""


def read_text(path: Path | str) -> str:
    """Return the document's text, dispatching on extension."""
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise UnsupportedDocument(
            f"no parser for {suffix!r}; supported: {sorted(SUPPORTED)}"
        )
    if suffix in {".md", ".markdown", ".txt"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return _read_xlsx(path)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    # Page markers survive into the text so the structurer can record page refs.
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        pages.append(f"<!--page:{i}-->\n{page.extract_text() or ''}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    out: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Recover heading level from the style so markdown structure survives.
        style = (para.style.name or "").lower()
        m = re.match(r"heading (\d)", style)
        out.append(f"{'#' * (int(m.group(1)) + 0)} {text}" if m else text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(out)


def _read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True)
    out: list[str] = []
    for sheet in wb.worksheets:
        out.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).strip() for c in row]
            if any(cells):
                out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


# --------------------------------------------------------------------------------------
# Shared text utilities
# --------------------------------------------------------------------------------------

_PAGE_MARKER = re.compile(r"<!--page:(\d+)-->")
_WS = re.compile(r"\s+")


def strip_page_markers(text: str) -> str:
    return _PAGE_MARKER.sub("", text)


def normalize(text: str) -> str:
    """Collapse whitespace and markdown emphasis for comparison purposes."""
    text = re.sub(r"[*_`]+", "", text)
    return _WS.sub(" ", text).strip()


def split_sentences(text: str) -> list[str]:
    """Cheap sentence splitter.

    Deliberately not a model. It protects the common abbreviations that appear in RFPs
    (No., Ltd., e.g., i.e., Rs.) and decimal numbers, which is enough for cue-word
    extraction and for per-sentence provenance later.
    """
    text = normalize(text)
    if not text:
        return []
    protected = text
    for abbr in ["No.", "Ltd.", "Pvt.", "Inc.", "e.g.", "i.e.", "etc.", "vs.", "Rs.",
                 "Dr.", "Mr.", "Mrs.", "Ms.", "Sr.", "Jr.", "Fig.", "Sec.", "Art."]:
        protected = protected.replace(abbr, abbr.replace(".", "\x00"))
    # NUL stands in for a dot that must not end a sentence. The replacement cannot be a
    # raw string: r"\x00" is a literal backslash-x, which re rejects as a bad escape.
    protected = re.sub(r"(\d)\.(\d)", "\\1\x00\\2", protected)

    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z(\[])", protected)
    return [p.replace("\x00", ".").strip() for p in parts if p.replace("\x00", ".").strip()]

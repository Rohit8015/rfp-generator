"""W3 Assembler — Phase 10. DETERMINISTIC: zero LLM calls.

Merges sections in outline order, injects charts, appends the compliance matrix,
renders Markdown -> docx (optional PDF), and emits the automation report:
% automated, per-section source breakdown, GAP list, consistency status.

The automation report is the project's headline number, so it is computed here from
provenance records rather than asserted. A section counts as automated only if no
sentence in it carries STAKEHOLDER provenance. Partial credit would let a section that
needed a human be reported as mostly automated, which is exactly the overclaim the
metric exists to prevent.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

from src.assurance.compliance import render_matrix
from src.models.schemas import (
    AssuranceFinding,
    AutomationReport,
    ComplianceMatrix,
    ConsistencyReport,
    GeneratedSection,
    ProvenanceKind,
    ResponseOutline,
    Severity,
)
from src.utils import metrics

log = logging.getLogger(__name__)


@dataclass
class Package:
    """Everything the run produced."""

    markdown_path: Path
    docx_path: Path | None
    report_path: Path
    report: AutomationReport
    section_count: int = 0
    asset_paths: list[str] = field(default_factory=list)


class Assembler:
    """Assembles and writes the final package. One public method: assemble()."""

    def __init__(self, settings=None) -> None:
        if settings is None:
            from config import get_settings

            settings = get_settings()
        self.settings = settings
        self.output_dir = Path(settings.output_path)

    # --- public ---------------------------------------------------------------------

    def assemble(
        self,
        run_id: str,
        outline: ResponseOutline,
        sections: list[GeneratedSection],
        matrix: ComplianceMatrix,
        consistency: ConsistencyReport,
        findings: list[AssuranceFinding],
        gap_requirement_ids: list[str],
        tasks_markdown: str = "",
        title: str = "Proposal response",
    ) -> Package:
        ordered = self._ordered(outline, sections)
        markdown = self._render(title, ordered, matrix, tasks_markdown, findings)

        self.output_dir.mkdir(parents=True, exist_ok=True)
        markdown_path = self.output_dir / f"{run_id}_response.md"
        markdown_path.write_text(markdown, encoding="utf-8")

        report = self.automation_report(run_id, ordered, matrix, consistency,
                                        gap_requirement_ids)
        report_path = self.output_dir / f"{run_id}_automation_report.md"
        report_path.write_text(self.render_report(report, findings, ordered),
                               encoding="utf-8")

        docx_path = self._write_docx(run_id, title, ordered, matrix, tasks_markdown,
                                     findings)

        return Package(
            markdown_path=markdown_path,
            docx_path=docx_path,
            report_path=report_path,
            report=report,
            section_count=len(ordered),
            asset_paths=[p for s in ordered for p in s.asset_paths],
        )

    # --- assembly -------------------------------------------------------------------

    @staticmethod
    def _ordered(outline: ResponseOutline, sections: list[GeneratedSection]
                 ) -> list[GeneratedSection]:
        """Outline order is the argument's order. Never emit in generation order."""
        by_id = {s.section_id: s for s in sections}
        ordered = [by_id[s.id] for s in outline.sections if s.id in by_id]
        # Anything generated but not in the outline still ships, at the end, visibly.
        extra = [s for s in sections if s.section_id not in {o.id for o in outline.sections}]
        return ordered + extra

    def _render(
        self,
        title: str,
        sections: list[GeneratedSection],
        matrix: ComplianceMatrix,
        tasks_markdown: str,
        findings: list[AssuranceFinding],
    ) -> str:
        stamp = datetime.now(timezone.utc).strftime("%d %B %Y")
        parts = [f"# {title}", "", f"Prepared {stamp}", "", "---", ""]

        for section in sections:
            parts.append(section.content_md.rstrip())
            for asset in section.asset_paths:
                name = Path(asset).name
                parts.append(f"\n![{section.title}](charts/{name})\n")
            parts.append("\n---\n")

        parts.append(render_matrix(matrix))
        if tasks_markdown:
            parts += ["\n---\n", tasks_markdown]

        blockers = [f for f in findings if f.severity is Severity.BLOCKER]
        if blockers:
            parts += ["\n---\n", "## Blocking assurance findings", "",
                      "These must be resolved before submission.", ""]
            parts += [f"- **{f.finding_type.value}** ({f.section_id or 'document'}): "
                      f"{f.detail}" for f in blockers]
            parts.append("")
        return "\n".join(parts)

    # --- the headline metric --------------------------------------------------------

    @staticmethod
    def automation_report(
        run_id: str,
        sections: list[GeneratedSection],
        matrix: ComplianceMatrix,
        consistency: ConsistencyReport,
        gap_requirement_ids: list[str],
    ) -> AutomationReport:
        return AutomationReport(
            run_id=run_id,
            overall_automation_rate=round(metrics.automation_rate(sections), 1),
            rate_by_form=metrics.automation_rate_by_form(sections),
            provenance_breakdown=metrics.provenance_breakdown(sections),
            gap_requirement_ids=list(gap_requirement_ids),
            consistency_passed=consistency.passed,
            compliance_coverage_pct=matrix.coverage_pct,
        )

    @staticmethod
    def render_report(report: AutomationReport, findings: list[AssuranceFinding],
                      sections: list[GeneratedSection]) -> str:
        automated = [s for s in sections if s.automated()]
        manual = [s for s in sections if not s.automated()]

        lines = [
            "# Automation report",
            "",
            f"Run `{report.run_id}`",
            "",
            "## Headline",
            "",
            "| Metric | Value |",
            "|---|---|",
            f"| Automation rate (sections) | **{report.overall_automation_rate:.1f}%** |",
            f"| Sections produced with no human input | {len(automated)} of {len(sections)} |",
            f"| Automation rate (sentences) | {metrics.sentence_automation_rate(sections):.1f}% |",
            f"| Requirement coverage | {report.compliance_coverage_pct:.1f}% |",
            f"| Consistency | {'passed' if report.consistency_passed else 'FAILED'} |",
            f"| Evidence gaps | {len(report.gap_requirement_ids)} |",
            "",
            "The section rate is the headline and is deliberately unforgiving: a section "
            "counts as automated only when no sentence in it needed a human, because a "
            "section with one carve-out is still a section someone must open. The "
            "sentence rate is reported alongside it because when most sections carry a "
            "small carve-out the section rate collapses to zero and stops distinguishing "
            "a document that is mostly drafted from one that is barely started.",
            "",
        ]

        if report.rate_by_form:
            lines += ["## By section type", "", "| Form | Automation rate |", "|---|---|"]
            lines += [f"| {form} | {rate:.1f}% |"
                      for form, rate in sorted(report.rate_by_form.items())]
            lines.append("")

        if report.provenance_breakdown:
            total = sum(report.provenance_breakdown.values()) or 1
            lines += ["## Where the text came from", "",
                      "| Provenance | Sentences | Share |", "|---|---|---|"]
            for kind, count in sorted(report.provenance_breakdown.items(),
                                      key=lambda t: -t[1]):
                lines.append(f"| {kind} | {count} | {100 * count / total:.1f}% |")
            lines.append("")

        if manual:
            lines += ["## Sections requiring human authorship", ""]
            lines += [f"- {s.section_id} {s.title} ({s.status.value})" for s in manual]
            lines.append("")

        if report.gap_requirement_ids:
            lines += ["## Evidence gaps", "",
                      "No proof point supports these requirements. They are surfaced, "
                      "never written around.", ""]
            lines += [f"- {rid}" for rid in report.gap_requirement_ids]
            lines.append("")

        by_type: dict[str, int] = {}
        for finding in findings:
            by_type[finding.finding_type.value] = by_type.get(finding.finding_type.value, 0) + 1
        if by_type:
            lines += ["## Assurance findings", "", "| Type | Count |", "|---|---|"]
            lines += [f"| {k} | {v} |" for k, v in sorted(by_type.items())]
            lines.append("")

        return "\n".join(lines)

    # --- docx -----------------------------------------------------------------------

    def _write_docx(self, run_id: str, title: str, sections: list[GeneratedSection],
                    matrix: ComplianceMatrix, tasks_markdown: str = "",
                    findings: list[AssuranceFinding] | None = None) -> Path | None:
        try:
            import docx
        except ImportError:  # pragma: no cover - python-docx is a hard dependency
            log.warning("python-docx unavailable; skipping docx export")
            return None

        document = docx.Document()
        document.add_heading(title, level=0)

        for section in sections:
            document.add_heading(section.title, level=1)
            for block in section.content_md.split("\n\n"):
                block = block.strip()
                if not block or block.startswith("#"):
                    continue
                if block.startswith("|"):
                    self._add_table(document, block)
                    continue
                document.add_paragraph(block)
            for asset in section.asset_paths:
                if Path(asset).is_file():
                    try:
                        document.add_picture(asset, width=docx.shared.Inches(6.0))
                    except Exception as exc:  # noqa: BLE001 - a bad image is not fatal
                        log.warning("could not embed %s: %s", asset, exc)

        document.add_heading("Requirements compliance matrix", level=1)
        self._add_matrix_table(document, matrix)

        # Human tasks and blocking findings, so the Word file carries the same
        # human-facing hand-off as the Markdown export rather than stopping at the matrix.
        self._add_tasks(document, tasks_markdown)
        self._add_blockers(document, findings or [])

        path = self.output_dir / f"{run_id}_response.docx"
        document.save(str(path))
        return path

    def _add_tasks(self, document, tasks_markdown: str) -> None:
        """Render the Human tasks section. Reuses the Markdown the tracker produced."""
        if not tasks_markdown.strip():
            return
        document.add_heading("Human tasks", level=1)
        for block in tasks_markdown.split("\n\n"):
            block = block.strip()
            if not block or block.startswith("#"):
                continue
            if block.startswith("|"):
                self._add_table(document, block)
            else:
                document.add_paragraph(block)

    @staticmethod
    def _add_blockers(document, findings: list[AssuranceFinding]) -> None:
        """List the findings that must be resolved before submission."""
        blockers = [f for f in findings if f.severity is Severity.BLOCKER]
        if not blockers:
            return
        document.add_heading("Blocking assurance findings", level=1)
        document.add_paragraph("These must be resolved before submission.")
        for finding in blockers:
            document.add_paragraph(
                f"{finding.finding_type.value} "
                f"({finding.section_id or 'document'}): {finding.detail}",
                style="List Bullet",
            )

    @staticmethod
    def _add_table(document, block: str) -> None:
        rows = [
            [c.strip() for c in line.strip().strip("|").split("|")]
            for line in block.splitlines()
            if line.strip().startswith("|") and not set(line.strip()) <= set("|-: ")
        ]
        if not rows:
            return
        width = max(len(r) for r in rows)
        table = document.add_table(rows=0, cols=width)
        table.style = "Light Grid Accent 1"
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row[:width]):
                cells[index].text = value.replace("**", "")

    @staticmethod
    def _add_matrix_table(document, matrix: ComplianceMatrix) -> None:
        table = document.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        for index, label in enumerate(
            ["Requirement", "Priority", "Section", "Anchor", "RAG"]
        ):
            header[index].text = label
        for row in matrix.rows:
            cells = table.add_row().cells
            cells[0].text = row.requirement_id
            cells[1].text = row.priority.value
            cells[2].text = row.section_id or "—"
            cells[3].text = row.anchor or "—"
            cells[4].text = row.rag.value


def provenance_kind_counts(sections: list[GeneratedSection]) -> dict[str, int]:
    """Convenience for the dashboard."""
    counts: dict[str, int] = {}
    for section in sections:
        for record in section.sentences:
            counts[record.kind.value] = counts.get(record.kind.value, 0) + 1
    for kind in ProvenanceKind:
        counts.setdefault(kind.value, 0)
    return counts
